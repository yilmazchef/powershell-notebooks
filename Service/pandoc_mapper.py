
import os
import collections.abc
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def img2base64(image_file, output_file):

    # need base 64
    import base64
    import sys
    # open the image
    image = open(image_file, 'rb')
    # read it
    image_read = image.read()
    # encode it as base 64
    # after python>=3.9, use `encodebytes` instead of `encodestring`
    image_64_encode = base64.encodestring(image_read) if sys.version_info < (3, 9) else base64.encodebytes(image_read)
    # convert the image base 64 to a string
    image_string = str(image_64_encode)
    # replace the newline characters
    image_string = image_string.replace("\\n", "")
    # replace the initial binary
    image_string = image_string.replace("b'", "")
    # replace the final question mark
    image_string = image_string.replace("'", "")
    # add the image tags
    image_string = '<p><img src="data:image/png;base64,' + image_string + '"></p>'
    # write it out
    image_result = open(output_file, 'w')
    image_result.write(image_string)


def ipynb2md(ipynb_file: str) -> str:
    cmdlet = f"jupyter nbconvert --to markdown \"{ipynb_file}\""
    os.system(cmdlet)
    return ipynb_file.replace(".ipynb", ".md")


def md2ipynb(md_file: str) -> str:

    ipynb_file = md_file.replace(".md", ".ipynb")
    cmdlet = f"pandoc \"{md_file}\" -o \"{ipynb_file}\""
    os.system(cmdlet)
    return ipynb_file


def md2odt(md_file):
    odt_file = md_file.replace(".md", "odt")
    cmdlet = f"pandoc -t odt \"{md_file}\" -o \"{odt_file}\""
    os.system(cmdlet)


def md2pptx(md_file: str) -> str:

    c = collections
    c.abc = collections.abc

    pptx_path = md_file.replace(".md", ".pptx")
    cmdlet = f"pandoc -V fontsize=12pt \"{md_file}\" -s --wrap auto -o \"{pptx_path}\""
    os.system(cmdlet)

    return pptx_path


def md2docx(md_file: str) -> str:

    docx_file = md_file.replace(".md", ".docx")
    cmdlet = f"pandoc \"{md_file}\" -f markdown -o \"{docx_file}\""
    os.system(cmdlet)
    return docx_file


def h4docx(docx_file, header_image, header_text=None):

    # checking if file already present and creating it if not present
    if os.path.isfile(rf"{docx_file}"):

        # opening the existing document
        document = Document(docx_file)

        header = document.sections[0].header

        htable = header.add_table(1, 2, Inches(4))
        htab_cells = htable.rows[0].cells
        ht0 = htab_cells[0].add_paragraph()
        kh = ht0.add_run()
        kh.add_picture(header_image, width=Inches(3))

        if header_text is not None:
            ht1 = htab_cells[1].add_paragraph(header_text)
            ht1.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # saving the blank document
        document.save(docx_file)


def f4docx(docx_file, footer_image, footer_text=None):

    document = None

    # checking if file already present and creating it if not present
    if os.path.isfile(rf"{docx_file}"):

        # opening the existing document
        document = Document(docx_file)

        footer = document.sections[0].footer

        ftable = footer.add_table(1, 2, Inches(4))
        ftab_cells = ftable.rows[0].cells
        ft0 = ftab_cells[0].add_paragraph()
        fh = ft0.add_run()

        imgSize = Inches(4)

        if footer_text is not None:
            imgSize = Inches(3)
            ft1 = ftab_cells[1].add_paragraph(footer_text)
            ft1.alignment = WD_ALIGN_PARAGRAPH.CENTER

        fh.add_picture(footer_image, width=imgSize)

        document.save(docx_file)


def md2pdf(md_file: str) -> str:

    pdf_file = md_file.replace(".md", ".pdf")
    cmdlet = f"pandoc \"{md_file}\" --pdf-engine=xelatex -o \"{pdf_file}\""
    os.system(cmdlet)
    return pdf_file
